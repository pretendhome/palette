#!/usr/bin/env python3
"""
Parameterized resume builder for Palette talent skill.

Usage:
  python3 build_resume.py --profile knowledge_data_engineer --output ~/Downloads/resume.docx
  python3 build_resume.py --profile certification_architect --output ~/Downloads/resume.docx
  python3 build_resume.py --profile technical_marketing_engineer --output ~/Downloads/resume.docx
  python3 build_resume.py --profile forward_deployed_engineer --output ~/Downloads/resume.docx
  python3 build_resume.py --profile custom --headline "..." --summary "..." --output ~/Downloads/resume.docx

The builder reads experience-inventory.yaml and assembles the resume from
structured career data. Each profile selects different headline, summary,
bullet variants, and Palette framing for the same underlying experience.
"""

import argparse
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).parent
INVENTORY_PATH = SCRIPT_DIR / "experience-inventory.yaml"

DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)
ACCENT = RGBColor(0x2C, 0x3E, 0x50)


def load_inventory():
    with open(INVENTORY_PATH) as f:
        return yaml.safe_load(f)


def set_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def add_text(paragraph, text, font_name="Arial", size=11, bold=False, italic=False, color=DARK):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '999999',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_capability_table(doc, capabilities):
    """Two-column borderless table of capabilities."""
    p = doc.add_paragraph()
    add_text(p, "CORE CAPABILITIES", size=10.5, bold=True, color=ACCENT)
    set_spacing(p, before=4, after=3)

    mid = (len(capabilities) + 1) // 2
    table = doc.add_table(rows=mid, cols=2)
    table.autofit = True
    for i in range(mid):
        for j in range(2):
            idx = i + j * mid
            if idx < len(capabilities):
                cell = table.cell(i, j)
                cell.text = ""
                cp = cell.paragraphs[0]
                add_text(cp, "\u2022  " + capabilities[idx], size=8.5, color=DARK)
                set_spacing(cp, before=1, after=1, line=11)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.makeelement(qn('w:tcMar'), {})
                for side in ['top', 'bottom', 'start', 'end']:
                    el = tcMar.makeelement(qn(f'w:{side}'), {qn('w:w'): '20', qn('w:type'): 'dxa'})
                    tcMar.append(el)
                tcPr.append(tcMar)

    # Remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.makeelement(qn(f'w:{bn}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(b)
    tblPr.append(borders)


# --- PROFILE DEFINITIONS ---
# Each profile defines how to frame the same career for a different role.

PROFILES = {
    "knowledge_data_engineer": {
        "headline": "AI KNOWLEDGE & DATA ENGINEER  |  KNOWLEDGE ARCHITECTURE & RETRIEVAL SYSTEMS",
        "summary": (
            "12+ years at Amazon building knowledge architectures, ontologies, and AI-augmented retrieval "
            "systems at enterprise scale. Designed and shipped a 117-node problem taxonomy with service routing, "
            "a 167-entry sourced knowledge library with evidence tiers and retrieval metadata, and a multi-agent "
            "orchestration system with 9 specialized agents, 1,876-quad relationship graph, and governed message "
            "bus. Background in comparative linguistics and knowledge engineering \u2014 the discipline of turning "
            "unstructured domain knowledge into machine-navigable structures."
        ),
        "bullet_key": "data_engineer",
        "capabilities": [
            "Knowledge architecture: ontology design, taxonomy construction, semantic data modeling, entity resolution",
            "Retrieval system design: structured knowledge retrieval, graph-based query, taxonomy-driven routing",
            "Data pipeline engineering: ingestion, transformation, indexing, and validation for heterogeneous sources",
            "LLM knowledge enrichment: RAG pipeline design, context assembly, calibration exemplars, prompt engineering",
            "Agent memory & orchestration: cognitive memory, multi-agent coordination, persistent context",
            "Knowledge graph construction: relationship extraction, quad-based graphs, cross-layer validation",
            "AI evaluation systems: rubric design, threshold engines, automated integrity checks, confidence scoring",
            "Python, YAML/structured data, SQL, Git/GitHub, Claude API, Claude Code, LLM evaluation frameworks",
        ],
        "palette_items": [
            ("Taxonomy & ontology", "117-node problem taxonomy (RIUs) with classification metadata, difficulty ratings, prerequisite DAG, and service routing"),
            ("Knowledge library", "167 sourced entries with evidence tiers (Tier 1: Google/Anthropic/OpenAI/AWS; Tier 2: NIST/peer-reviewed; Tier 3: validated open-source), journey stages, and cross-referencing to taxonomy nodes"),
            ("Knowledge graph", "1,876-quad relationship graph enabling bidirectional traversal across taxonomy, knowledge, services, people signals, and integration recipes"),
            ("Multi-agent orchestration", "9 specialized agents with cognitive memory, governed message bus (MCP protocol), and persistent context across interactions"),
            ("Data pipelines", "14 automated integrity validators, cross-layer consistency verification, threshold engine (89/89 tests), and continuous validation"),
            ("Retrieval infrastructure", "Taxonomy-driven routing classifying 117 problem types as internal_only (80) or service-routable (37), with 69 integration recipes"),
        ],
        "palette_tools": "Python, Claude API, Claude Code, YAML-based knowledge representation, SQL, Git/GitHub",
    },
    "certification_architect": {
        "headline": "CERTIFICATION & KNOWLEDGE SYSTEMS ARCHITECT  |  AI-AUGMENTED EDUCATION DESIGN",
        "summary": (
            "12+ years at Amazon building knowledge architectures, competency frameworks, and AI-augmented "
            "assessment systems at enterprise scale. Designed and shipped a developer certification system with "
            "117 curriculum modules, 5 certification journeys, portfolio-based assessment with a 3-layer evaluation "
            "pipeline (automated validators \u2192 AI rubric scoring with calibration exemplars \u2192 human review), "
            "and a threshold engine enforcing competency gates across all modules. Skeptical of multiple-choice "
            "as the default \u2014 built assessment systems where AI evaluates demonstrations of competence, not "
            "recall. Deep hands-on fluency in Claude\u2019s products (API, Claude Code, Claude.ai)."
        ),
        "bullet_key": "enablement",
        "capabilities": [
            "Certification architecture: curriculum design, competency frameworks, learning progressions, credentialing systems",
            "AI-augmented assessment: 3-layer evaluation pipeline, LLM-as-judge scoring, calibration exemplar design",
            "Item bank design: calibration exemplars differentiating by quality of thinking, rubric development, threshold policy",
            "Portfolio-based assessment: performance evaluation, artifact-based competency measurement, capstone project design",
            "Knowledge engineering: ontology design, taxonomy construction, structured knowledge at scale",
            "Intelligent tooling: building systems that maintain and evolve content through AI automation",
            "Cross-functional collaboration: GTM, Partnerships, Engineering, Product, Customer Success",
            "Python, Claude API, Claude Code, LLM evaluation frameworks, SQL, Git/GitHub",
        ],
        "palette_items": [
            ("Certification system", "117 curriculum modules across 6 workstreams, 5 certification journeys (AI Foundations, RAG Engineer, Agent Architect, AI Operations, AI Governance), 5 capstone projects with artifact-based evaluation"),
            ("3-layer assessment pipeline", "Automated validators \u2192 AI rubric scoring with calibration exemplars \u2192 human review escalation. Confidence-based: AI handles consistency, humans handle edge cases and novel reasoning"),
            ("Item banks & calibration", "11 calibration exemplar sets differentiating competency levels by quality of thinking, not word count. Threshold engine enforcing per-dimension competence on critical control modules"),
            ("Certification tiers", "UNVALIDATED \u2192 WORKING \u2192 PRODUCTION with rising rigor. PRODUCTION requires one-way-door decision defense, auditability, and mandatory human sign-off"),
            ("Content currency & renewal", "Automated staleness detection flags modules needing update as AI capabilities evolve; versioned curriculum with integrity validators ensuring cross-module consistency"),
        ],
        "palette_tools": "Python, Claude API, Claude Code, Claude.ai (daily usage), Git/GitHub",
    },
    "technical_marketing_engineer": {
        "headline": "TECHNICAL MARKETING ENGINEER  |  KNOWLEDGE ARCHITECTURE & DEVELOPER ENABLEMENT",
        "summary": (
            "12+ years at Amazon translating complex AI systems into adoption-driving content and developer "
            "enablement programs. Built Ask Pathfinder, a knowledge retrieval system serving 12,000+ sellers monthly "
            "with 25% reduction in prep time and 67% increase in sales-play coverage. Created the AWS Data "
            "Leadership Forum reaching 291 senior data leaders and 98 CxOs. Organized enterprise enablement programs "
            "featuring Mistral and other frontier AI partners. Background in comparative linguistics \u2014 the discipline "
            "of understanding how the same concept, expressed differently, changes whether people understand it, "
            "believe it, and act on it."
        ),
        "bullet_key": "tme",
        "capabilities": [
            "Technical content strategy: whitepapers, solution briefs, architecture guides, developer documentation",
            "Developer enablement: workshops, training programs, hands-on labs, adoption measurement",
            "Executive engagement: C-level content, decision frameworks, technical community building",
            "Product positioning: translating model capabilities into enterprise deployment narratives",
            "Partner enablement: frontier AI company partnerships, ecosystem technical content",
            "Knowledge architecture: taxonomy design, metadata frameworks, retrieval system design",
            "Adoption measurement: three-level instrumentation (usage \u2192 behavior shift \u2192 business outcomes)",
            "Python, Claude API, Claude Code, SQL, Git/GitHub, presentation design",
        ],
        "palette_items": [
            ("Knowledge retrieval", "Ask Pathfinder-inspired 3-layer architecture: content ingestion, metadata layer with human-designed signals, use-case ontology spanning 15 verticals"),
            ("Content taxonomy", "117-node problem taxonomy enabling content discovery by problem type, not product name"),
            ("Enablement system", "117 curriculum modules across 6 workstreams with competency graph, 5 certification journeys, 3-layer assessment pipeline (automated validators \u2192 AI rubric scoring with calibration exemplars \u2192 human review), and threshold engine enforcing per-dimension competence"),
            ("Multi-agent collaboration", "Coordinated enablement build across Claude, Codex, Kiro, and Mistral relay patterns \u2014 each agent owning a workstream with structured handoffs and semantic integrity validation"),
            ("Adoption instrumentation", "Three-level measurement: usage metrics \u2192 behavior shift indicators \u2192 business outcome correlation"),
        ],
        "palette_tools": "Python, Claude API, Claude Code, Git/GitHub",
    },
    "forward_deployed_engineer": {
        "headline": "FORWARD DEPLOYED ENGINEER  |  AI SYSTEMS ARCHITECTURE & CUSTOMER DELIVERY",
        "summary": (
            "12+ years at Amazon building and deploying AI knowledge systems at enterprise scale \u2014 from a "
            "25-billion-node Knowledge Graph at Alexa Automotive to a knowledge retrieval system serving 12,000+ "
            "AWS sellers monthly. Hands-on builder: designed entity resolution pipelines, Spark ETL across 47 data "
            "providers, vector-based deduplication, structured attribution for hallucination detection, and a "
            "multi-agent orchestration system with 9 specialized agents. Background in comparative linguistics and "
            "knowledge engineering, with deep fluency across Python, Claude API, and production AI systems."
        ),
        "bullet_key": "data_engineer",
        "capabilities": [
            "AI systems architecture: end-to-end design from data ingestion to production deployment",
            "Customer-facing technical delivery: scoping, architecture review, integration support, success measurement",
            "Knowledge engineering: ontology design, entity resolution, knowledge graph construction, retrieval systems",
            "Data pipeline engineering: Spark ETL, multi-source ingestion, vector embeddings, quality validation",
            "LLM systems: RAG pipeline design, hallucination detection, structured attribution, prompt engineering",
            "Multi-agent orchestration: cognitive memory, governed message bus, persistent context across interactions",
            "Cross-functional execution: partnering with Product, Engineering, GTM, and Customer Success",
            "Python, Go, SQL, Spark, Claude API, Claude Code, Git/GitHub",
        ],
        "palette_items": [
            ("Knowledge graph", "1,876-quad relationship graph with bidirectional traversal across taxonomy, knowledge, services, people signals, and integration recipes"),
            ("Multi-agent system", "9 specialized agents with cognitive memory, governed message bus (MCP protocol), and persistent context"),
            ("Data pipelines", "14 automated integrity validators, cross-layer consistency verification, threshold engine (89/89 tests)"),
            ("Retrieval infrastructure", "Taxonomy-driven routing: 80 internal-only + 37 service-routable problem types, 69 integration recipes"),
        ],
        "palette_tools": "Python, Go, Claude API, Claude Code, Spark, SQL, Git/GitHub",
    },
    "enablement_strategy": {
        "headline": "AI ENABLEMENT & STRATEGY LEAD  |  SCALED ADOPTION & EXECUTIVE ENGAGEMENT",
        "summary": (
            "12+ years at Amazon driving AI adoption at enterprise scale \u2014 from building Ask Pathfinder, "
            "a knowledge retrieval system serving 12,000+ sellers monthly with 25% reduction in prep time, "
            "to creating the AWS Data Leadership Forum reaching 291 senior data leaders and 98 CxOs. "
            "Delivered 250+ technical enablement sessions reaching 20,000+ users annually, driving +17% "
            "engagement and +67% feature adoption. Measures at three levels: usage, behavior shift, and "
            "business outcomes. Background in comparative linguistics \u2014 the discipline of understanding "
            "how the same concept, expressed differently, changes whether people act on it."
        ),
        "bullet_key": "enablement",
        "capabilities": [
            "AI adoption strategy: scaled enablement programs, change management, adoption measurement",
            "Executive engagement: C-level content, decision frameworks, technical community building (291+ leaders)",
            "Measurement design: three-level instrumentation (usage \u2192 behavior shift \u2192 business outcomes)",
            "Knowledge architecture: taxonomy design, metadata frameworks, retrieval system design",
            "Cross-functional leadership: GTM, Product, Engineering, Partnerships, Customer Success",
            "Partner enablement: frontier AI company partnerships within 140,000-person ecosystem",
            "Technical content: workshops, training programs, enablement materials at scale",
            "Python, Claude API, Claude Code, SQL, Git/GitHub",
        ],
        "palette_items": [
            ("Enablement taxonomy", "117-node problem taxonomy mapping enterprise AI challenges to adoption pathways"),
            ("Knowledge library", "167 sourced entries with evidence tiers and learning progressions (foundation \u2192 specialization)"),
            ("Adoption measurement", "Three-level instrumentation: usage metrics, behavior shift indicators, business outcome correlation"),
            ("Operational scaling", "Intelligent automation maintaining content currency without headcount growth"),
        ],
        "palette_tools": "Python, Claude API, Claude Code, Git/GitHub",
    },
    "customer_success_ai": {
        "headline": "AI DEPLOYMENT & CUSTOMER SUCCESS  |  OUTCOME-DRIVEN AI ADOPTION",
        "summary": (
            "12+ years at Amazon building and deploying AI systems that drive measurable customer outcomes. "
            "Built Ask Pathfinder, a knowledge retrieval system serving 12,000+ sellers monthly with 25% reduction "
            "in prep time and +67% feature adoption. Designed evaluation frameworks for AI output quality at "
            "Amazon AGI \u2014 structured attribution pipelines turning subjective quality assessment into measurable "
            "claim traceability. Organized enterprise enablement programs featuring Mistral and other frontier AI "
            "partners, translating model capabilities into enterprise deployment narratives. Background in comparative linguistics and "
            "knowledge engineering."
        ),
        "bullet_key": "enablement",
        "capabilities": [
            "Customer outcomes: deployment strategy, adoption measurement, expansion planning, time-to-value optimization",
            "AI evaluation: quality frameworks, structured rubrics, confidence scoring, hallucination detection",
            "Adoption measurement: three-level instrumentation (usage \u2192 behavior shift \u2192 business outcomes)",
            "Technical delivery: 250+ enablement sessions, 20,000+ users reached, partner ecosystem enablement",
            "Executive engagement: 291+ senior data leaders, 98+ CxOs, decision-framework content",
            "Partner management: frontier AI companies (Mistral, xAI, Stability AI, Luma AI, TwelveLabs)",
            "Knowledge architecture: taxonomy design, retrieval systems, data quality instrumentation",
            "Python, Claude API, Claude Code, SQL, Git/GitHub",
        ],
        "palette_items": [
            ("Customer intelligence", "Taxonomy-driven routing classifying 117 problem types with service recommendations"),
            ("Quality instrumentation", "14 automated integrity validators, threshold engine, cross-layer consistency verification"),
            ("Evaluation system", "3-layer assessment (automated checks \u2192 AI rubric \u2192 human review) with calibration exemplars"),
            ("Knowledge retrieval", "167 sourced entries with evidence tiers enabling precision knowledge delivery"),
        ],
        "palette_tools": "Python, Claude API, Claude Code, Git/GitHub",
    },
    "learning_systems_engineer": {
        "headline": "LEARNING SYSTEMS ENGINEER  |  AI-NATIVE EDUCATION INFRASTRUCTURE",
        "summary": (
            "12+ years at Amazon building knowledge architectures, learning infrastructure, and AI-augmented "
            "assessment systems at enterprise scale. Designed and shipped a production learning system with "
            "117 curriculum modules, 5 certification journeys, adaptive learning progressions "
            "(foundation \u2192 retrieval \u2192 orchestration \u2192 specialization), and a 3-layer evaluation pipeline "
            "(automated validators \u2192 AI rubric scoring with calibration exemplars \u2192 human review). Built "
            "analytics measuring learning outcomes \u2014 not vanity metrics: competency progression, inter-rater "
            "reliability (Cohen\u2019s kappa), and per-dimension threshold enforcement. Background in comparative "
            "linguistics and knowledge engineering \u2014 translating pedagogical goals into production systems."
        ),
        "bullet_key": "enablement",
        "capabilities": [
            "Learning infrastructure: curriculum architecture, adaptive progressions, prerequisite DAGs, staleness detection",
            "Assessment systems: 3-layer evaluation pipeline, LLM-as-judge, calibration exemplar design, threshold engines",
            "Learning analytics: competency tracking, inter-rater reliability (Cohen\u2019s kappa), coverage metrics, outcome measurement",
            "Credentialing & certification: competency frameworks, certification tiers, portfolio-based assessment, capstone design",
            "Knowledge engineering: ontology design, taxonomy construction, structured knowledge at scale (167 entries, 542 sources)",
            "Educator empowerment: non-technical onboarding systems, guided learning paths, plain-language instructional design",
            "AI-native workflows: LLM evaluation, confidence scoring, automated content validation, multi-agent coordination",
            "Python, Claude API, Claude Code, data pipelines, SQL, Git/GitHub",
        ],
        "palette_items": [
            ("Learning infrastructure", "117 curriculum modules across 6 workstreams, 5 certification journeys (AI Foundations, RAG Engineer, Agent Architect, AI Operations, AI Governance), prerequisite DAG (valid acyclic, max depth 3), automated integrity validation"),
            ("Adaptive learning system", "Learning progressions (foundation \u2192 retrieval \u2192 orchestration \u2192 specialization), certification tiers (UNVALIDATED \u2192 WORKING \u2192 PRODUCTION) with rising rigor, staleness detection flagging modules as AI capabilities evolve"),
            ("Assessment pipeline", "3-layer evaluation: automated validators \u2192 AI rubric scoring with calibration exemplars and confidence reporting \u2192 human review escalation. 11 calibration exemplar sets differentiating by quality of thinking"),
            ("Learning analytics", "Coverage reports (134/167 KL utilization = 80.2%), per-dimension competency tracking, threshold engine enforcing mastery on critical modules, disagreement analysis for rubric refinement"),
            ("Educator tools", "Enablement coach system for non-technical learners: 7-stage guided path, session-based progression, plain-language principles. Elia onboarding: 1-hour guided session plan"),
            ("Content currency & renewal", "Automated staleness detection, versioned curriculum with integrity validators (117/117 pass), cross-module consistency checks"),
        ],
        "palette_tools": "Python, Claude API, Claude Code, Claude.ai (daily usage), Git/GitHub",
    },
}


def build_resume(profile_name, output_path, inventory, custom_headline=None, custom_summary=None):
    profile = PROFILES.get(profile_name)
    if not profile and profile_name != "custom":
        print(f"Unknown profile: {profile_name}")
        print(f"Available: {', '.join(PROFILES.keys())}, custom")
        return

    if profile_name == "custom":
        if not custom_headline or not custom_summary:
            print("Custom profile requires --headline and --summary arguments")
            return
        # Use knowledge_data_engineer as base template for capabilities and palette items
        profile = PROFILES["knowledge_data_engineer"].copy()

    headline = custom_headline or profile["headline"]
    summary = custom_summary or profile["summary"]
    bullet_key = profile.get("bullet_key", "knowledge_engineer")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    prof = inventory["profile"]

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, prof["name"].upper(), size=20, bold=True, color=ACCENT)
    set_spacing(p, before=0, after=2)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    github = prof.get('portfolio', '')
    contact = f"{prof['location']}  \u2022  {prof['phone']}  \u2022  {prof['email']}  \u2022  {prof['linkedin']}"
    if github:
        contact += f" \u2022  {github}"
    add_text(p, contact, size=9, color=GRAY)
    set_spacing(p, before=0, after=4)

    add_horizontal_line(doc)

    # Headline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, headline, size=10.5, bold=True, color=ACCENT)
    set_spacing(p, before=4, after=4)

    # Summary
    p = doc.add_paragraph()
    add_text(p, summary, size=9.5, color=DARK)
    set_spacing(p, before=0, after=4)

    add_horizontal_line(doc)

    # Capabilities
    add_capability_table(doc, profile["capabilities"])
    add_horizontal_line(doc)

    # Professional Experience
    p = doc.add_paragraph()
    add_text(p, "PROFESSIONAL EXPERIENCE", size=10.5, bold=True, color=ACCENT)
    set_spacing(p, before=4, after=3)

    for era in reversed(inventory["eras"]):
        if era["id"] == "ERA-5":
            continue  # Palette goes in its own section

        # Company line
        p = doc.add_paragraph()
        add_text(p, era["company"], size=9.5, bold=True, color=DARK)
        add_text(p, f"  |  {era['location']}  |  {era['dates']}", size=9, color=GRAY)
        set_spacing(p, before=4, after=0)

        # Title
        title = era.get("headline_variants", {}).get(bullet_key, era["title"])
        p = doc.add_paragraph()
        add_text(p, title, size=9, bold=True, italic=True, color=DARK)
        set_spacing(p, before=0, after=2)

        # Bullets
        bullets = era.get("bullets", {}).get(bullet_key, era.get("bullets", {}).get("knowledge_engineer", []))
        for bullet in bullets:
            p = doc.add_paragraph()
            add_text(p, "\u2022  ", size=8.5, color=DARK)
            add_text(p, bullet, size=8.5, color=DARK)
            set_spacing(p, before=0, after=1, line=11)

    add_horizontal_line(doc)

    # Palette section
    p = doc.add_paragraph()
    add_text(p, "INDEPENDENT PROJECT: PALETTE INTELLIGENCE SYSTEM", size=10.5, bold=True, color=ACCENT)
    set_spacing(p, before=4, after=2)

    p = doc.add_paragraph()
    add_text(p, "Built a production AI knowledge architecture demonstrating end-to-end knowledge engineering at scale:", size=9, color=DARK)
    set_spacing(p, before=0, after=2)

    for label, desc in profile["palette_items"]:
        p = doc.add_paragraph()
        add_text(p, "\u2022  ", size=8.5, color=DARK)
        add_text(p, f"{label}: ", size=8.5, bold=True, color=DARK)
        add_text(p, desc, size=8.5, color=DARK)
        set_spacing(p, before=0, after=1, line=11)

    p = doc.add_paragraph()
    add_text(p, "\u2022  ", size=8.5, color=DARK)
    add_text(p, "Tools: ", size=8.5, bold=True, color=DARK)
    add_text(p, profile["palette_tools"], size=8.5, color=DARK)
    set_spacing(p, before=0, after=1, line=11)

    p = doc.add_paragraph()
    add_text(p, "\u2022  ", size=8.5, color=DARK)
    add_text(p, "Open source: ", size=8.5, bold=True, color=DARK)
    add_text(p, prof["portfolio"], size=8.5, color=DARK)
    set_spacing(p, before=0, after=1, line=11)

    add_horizontal_line(doc)

    # Education
    p = doc.add_paragraph()
    add_text(p, "EDUCATION", size=10.5, bold=True, color=ACCENT)
    set_spacing(p, before=4, after=2)

    for edu in inventory["education"]:
        p = doc.add_paragraph()
        add_text(p, f"{edu['school']} \u2014 {edu['degree']} ({edu['year']})", size=9, color=DARK)
        set_spacing(p, before=0, after=1)

    # Languages
    p = doc.add_paragraph()
    add_text(p, "LANGUAGES", size=10.5, bold=True, color=ACCENT)
    set_spacing(p, before=4, after=2)

    lang_str = ", ".join(
        f"{l['lang']} ({l['level']}" + (f" \u2014 {l['note']}" if l.get('note') else "") + ")"
        for l in prof["languages"]
    )
    p = doc.add_paragraph()
    add_text(p, lang_str, size=9, color=DARK)
    set_spacing(p, before=0, after=0)

    # Save
    doc.save(output_path)
    print(f"Resume saved to: {output_path}")

    # Verify
    print("\n--- Content Verification ---")
    verify = Document(output_path)
    for i, para in enumerate(verify.paragraphs):
        text = para.text.strip()
        if text:
            print(f"  [{i:3d}] {text[:120]}{'...' if len(text) > 120 else ''}")
    print(f"\nTotal paragraphs: {len(verify.paragraphs)}")


def main():
    parser = argparse.ArgumentParser(description="Build tailored resume from experience inventory")
    parser.add_argument("--profile", required=True, help="Role profile: " + ", ".join(PROFILES.keys()) + ", custom")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--headline", help="Custom headline (for --profile custom)")
    parser.add_argument("--summary", help="Custom summary (for --profile custom)")
    args = parser.parse_args()

    inventory = load_inventory()
    build_resume(args.profile, args.output, inventory, args.headline, args.summary)


if __name__ == "__main__":
    main()
